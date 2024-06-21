from docx import Document

def create_document():
    # Crear un nuevo documento
    doc = Document()

    # Agregar título al documento
    doc.add_heading('Secuenciación de Actividades', level=1)

    # Función para agregar subetapas y sus tablas de actividades
    def add_subetapa(doc, subetapa_title, data):
        doc.add_heading(subetapa_title, level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Subetapa'
        hdr_cells[1].text = 'Código de la actividad'
        hdr_cells[2].text = 'Descripción de la actividad'
        hdr_cells[3].text = 'Código del requisito'

        for item in data:
            row_cells = table.add_row().cells
            row_cells[0].text = item[0]
            row_cells[1].text = item[1]
            row_cells[2].text = item[2]
            row_cells[3].text = item[3]

    # Datos de cada subetapa
    data1 = [
        ["Planificación Inicial", "ACT01", "Reunión inicial con stakeholders", "RDC01"],
        ["Planificación Inicial", "ACT02", "Definición del equipo de proyecto", "RDC02"],
        ["Planificación Inicial", "ACT03", "Desarrollo del Acta de Constitución del Proyecto", "RDC03"],
        ["Planificación Inicial", "ACT04", "Reunión con stakeholders para recopilar requisitos", "RDC04"],
        ["Planificación Inicial", "ACT05", "Documentación de requisitos funcionales y no funcionales", "RDC05"],
        ["Planificación Inicial", "ACT06", "Evaluación de sistemas existentes", "RDC06"]
    ]

    data2 = [
        ["Diseño del Centro de Datos", "ACT07", "Diseño del layout físico de la sala del centro de datos", "RDC07"],
        ["Diseño del Centro de Datos", "ACT08", "Diseño del sistema HVAC", "RDC08"],
        ["Diseño del Centro de Datos", "ACT09", "Selección y adquisición de unidades de aire acondicionado y sistemas de ventilación", "RDC09"],
        ["Diseño del Centro de Datos", "ACT10", "Diseño de la infraestructura física", "RDC10"],
        ["Diseño del Centro de Datos", "ACT11", "Planificación del espacio físico", "RDC11"]
    ]

    data3 = [
        ["Seguridad y Respaldo", "ACT12", "Realizar una evaluación de riesgos y vulnerabilidades", "RDC12"],
        ["Seguridad y Respaldo", "ACT13", "Definir la estrategia de respaldo de datos", "RDC13"],
        ["Seguridad y Respaldo", "ACT14", "Selección y configuración de sistemas de almacenamiento para backups", "RDC13"],
        ["Seguridad y Respaldo", "ACT15", "Implementación de medidas de seguridad física y lógica", "RDC12"],
        ["Seguridad y Respaldo", "ACT16", "Implementación de sistemas de monitoreo y detección de intrusos", "RDC12"]
    ]

    data4 = [
        ["Desarrollo y Pruebas", "ACT17", "Seleccionar y adquirir servidores de alto rendimiento", "RDC09"],
        ["Desarrollo y Pruebas", "ACT18", "Configurar y desplegar sistemas de almacenamiento", "RDC09"],
        ["Desarrollo y Pruebas", "ACT19", "Diseñar la arquitectura de red interna y conexiones externas", "RDC10"],
        ["Desarrollo y Pruebas", "ACT20", "Instalar y configurar equipos de red", "RDC10"],
        ["Desarrollo y Pruebas", "ACT21", "Realizar pruebas de rendimiento y ajustes", "RDC09"]
    ]

    # Agregar las subetapas al documento
    add_subetapa(doc, 'Subetapa 1: Planificación Inicial', data1)
    add_subetapa(doc, 'Subetapa 2: Diseño del Centro de Datos', data2)
    add_subetapa(doc, 'Subetapa 3: Implementación de Seguridad y Respaldo', data3)
    add_subetapa(doc, 'Subetapa 4: Desarrollo y Pruebas', data4)

    # Guardar el documento
    try:
        doc.save(r"Secuenciacion_de_Actividades.docx")
        print("Documento creado con éxito")
    except Exception as e:
        print(f"Error al guardar el documento: {e}")

if __name__ == "__main__":
    create_document()
